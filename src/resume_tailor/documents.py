from __future__ import annotations

import json
import math
import re
from collections import Counter
from pathlib import Path
from typing import Any, Iterator, List, Tuple

import fitz
from docx import Document
from docx.text.paragraph import Paragraph

from .models import (
    ResumeBlock,
    ResumeStructure,
    StructuredResumeContent,
    TailoredBlock,
)


def _infer_block_type(text: str) -> str:
    t = text.strip()
    if not t:
        return "normal"
    if t.startswith(("•", "-", "*", "·")) or re.match(r"^\d+[\.)]\s", t):
        return "bullet"
    if len(t) < 80 and t.isupper():
        return "heading"
    if len(t) < 50 and not t.endswith(".") and t[0:1].isalpha() and t == t.title():
        return "heading"
    return "normal"


def _iter_docx_paragraphs(doc: Document) -> Iterator[Paragraph]:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def extract_resume_structure(path: str) -> ResumeStructure:
    p = Path(path).expanduser().resolve()
    if not p.is_file():
        raise FileNotFoundError(f"Resume not found: {p}")

    suffix = p.suffix.lower()
    if suffix == ".docx":
        blocks = _extract_docx_blocks(str(p))
    elif suffix == ".pdf":
        blocks = _extract_pdf_blocks(str(p))
    elif suffix == ".txt":
        blocks = _extract_txt_blocks(str(p))
    else:
        raise ValueError(f"Unsupported resume format: {suffix} (use .docx, .pdf, or .txt)")

    full = "\n".join(b.text for b in blocks if b.text.strip())
    return ResumeStructure(
        source_path=str(p),
        source_format=suffix.lstrip("."),
        blocks=blocks,
        full_text=full,
    )


def _extract_docx_blocks(path: str) -> List[ResumeBlock]:
    doc = Document(path)
    blocks: List[ResumeBlock] = []
    idx = 0
    for para in _iter_docx_paragraphs(doc):
        text = para.text
        style_name = para.style.name if para.style else ""
        bt = "heading" if "Heading" in style_name else _infer_block_type(text)
        blocks.append(ResumeBlock(index=idx, text=text, block_type=bt))
        idx += 1
    return blocks


def _extract_pdf_blocks(path: str) -> List[ResumeBlock]:
    doc = fitz.open(path)
    rows: List[Tuple[float, float, str]] = []
    try:
        for page in doc:
            for b in page.get_text("blocks"):
                x0, y0, x1, y1, text, *_ = b
                t = (text or "").strip()
                if t:
                    rows.append((y0, x0, t))
    finally:
        doc.close()
    if not rows:
        from pdfminer.high_level import extract_text

        fallback = extract_text(path)
        lines = [ln for ln in fallback.splitlines()]
        return [
            ResumeBlock(index=i, text=ln, block_type=_infer_block_type(ln))
            for i, ln in enumerate(lines)
        ]
    rows.sort(key=lambda r: (round(r[0], 1), r[1]))
    blocks: List[ResumeBlock] = []
    for i, (_, _, text) in enumerate(rows):
        blocks.append(ResumeBlock(index=i, text=text, block_type=_infer_block_type(text)))
    return blocks


def _extract_txt_blocks(path: str) -> List[ResumeBlock]:
    raw = Path(path).read_text(encoding="utf-8", errors="replace")
    lines = raw.splitlines()
    blocks: List[ResumeBlock] = []
    idx = 0
    for line in lines:
        blocks.append(ResumeBlock(index=idx, text=line, block_type=_infer_block_type(line)))
        idx += 1
    return blocks


def read_text_file_strict(path: str) -> str:
    """Read a UTF-8 text file; resume and job description inputs must be .txt in the template flow."""
    p = Path(path).expanduser().resolve()
    if not p.is_file():
        raise FileNotFoundError(f"File not found: {p}")
    if p.suffix.lower() != ".txt":
        raise ValueError(
            f"Only .txt is supported for this path (got {p.suffix}). Convert documents to plain text."
        )
    return p.read_text(encoding="utf-8", errors="replace")


def extract_plain_text(path: str) -> str:
    p = Path(path).expanduser().resolve()
    if not p.is_file():
        raise FileNotFoundError(f"File not found: {p}")

    suffix = p.suffix.lower()
    if suffix == ".txt":
        return p.read_text(encoding="utf-8", errors="replace")
    if suffix == ".docx":
        doc = Document(str(p))
        parts = [para.text for para in _iter_docx_paragraphs(doc)]
        return "\n".join(parts)
    if suffix == ".pdf":
        doc = fitz.open(str(p))
        try:
            return "\n".join(page.get_text() for page in doc)
        finally:
            doc.close()
    raise ValueError(f"Unsupported format for job description: {suffix}")


def apply_tailored_blocks_safe(
    original: ResumeStructure,
    tailored: List[TailoredBlock],
    docx_out: Path,
) -> None:
    by_index = {b.index: b.text for b in tailored}
    n = len(original.blocks)
    if len(by_index) != n:
        raise ValueError(
            f"Expected {n} tailored blocks (one per resume block), got {len(by_index)}"
        )
    for i in range(n):
        if i not in by_index:
            raise ValueError(f"Missing tailored text for block index {i}")

    src_path = original.source_path
    suffix = Path(src_path).suffix.lower()

    if suffix == ".docx":
        doc = Document(src_path)
        idx = 0
        for para in _iter_docx_paragraphs(doc):
            para.text = by_index[idx]
            idx += 1
        if idx != n:
            raise RuntimeError("DOCX paragraph count changed vs extraction")
        doc.save(str(docx_out))
        return

    doc = Document()
    for i in range(n):
        doc.add_paragraph(by_index[i])
    doc.save(str(docx_out))


def _docx_text_to_reportlab_markup(text: str) -> str:
    """Turn docx paragraph text into ReportLab Paragraph markup (XML subset)."""
    t = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    # Model output sometimes contains literal two-char sequences
    t = t.replace("\\n", "\n").replace("\\t", "\t")
    # Drop other C0 controls (invalid in JSON and ugly in PDF)
    t = "".join(ch if (ch >= " " or ch in "\n\t") else " " for ch in t)
    lines = t.split("\n")
    escaped: List[str] = []
    for line in lines:
        line = (
            line.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
        )
        line = line.replace("\t", "    ")
        escaped.append(line)
    return "<br/>".join(escaped)


def write_resume_pdf_from_docx(docx_path: Path, pdf_path: Path) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Paragraph as RLParagraph
    from reportlab.platypus import SimpleDocTemplate, Spacer

    docx_path = docx_path.resolve()
    d = Document(str(docx_path))
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "ResumeBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=13,
        spaceAfter=5,
    )
    head = ParagraphStyle(
        "ResumeSection",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        spaceAfter=6,
        spaceBefore=8,
    )
    title = ParagraphStyle(
        "ResumeTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=19,
        spaceAfter=10,
        spaceBefore=0,
    )
    h1 = ParagraphStyle(
        "ResumeH1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        spaceAfter=7,
        spaceBefore=6,
    )

    story = []
    for para in _iter_docx_paragraphs(d):
        raw = para.text or ""
        if not raw.strip():
            story.append(Spacer(1, 4))
            continue
        style_name = para.style.name if para.style else ""
        sn = style_name or ""
        if sn == "Title" or "Title" in sn:
            st = title
        elif sn.strip() in ("Heading 1",) or sn.startswith("Heading 1"):
            st = h1
        elif "Heading" in sn:
            st = head
        else:
            st = body
        markup = _docx_text_to_reportlab_markup(raw)
        story.append(RLParagraph(markup, st))

    pdf = SimpleDocTemplate(
        str(pdf_path),
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54,
        title="Resume",
    )
    pdf.build(story)


def simple_cosine_similarity(a: str, b: str) -> float:
    aw = re.findall(r"\w+", a.lower())
    bw = re.findall(r"\w+", b.lower())
    if not aw or not bw:
        return 0.0
    ca, cb = Counter(aw), Counter(bw)
    vocab = set(ca) | set(cb)
    va = [ca[w] for w in vocab]
    vb = [cb[w] for w in vocab]
    dot = sum(x * y for x, y in zip(va, vb))
    na = math.sqrt(sum(x * x for x in va))
    nb = math.sqrt(sum(y * y for y in vb))
    if na == 0 or nb == 0:
        return 0.0
    return dot / (na * nb)


def _loads_tailoring_payload(s: str):
    """Parse JSON from the LLM; repair common issues (raw newlines in strings, missing commas)."""
    raw = (s or "").strip()
    if not raw:
        raise ValueError("tailoring_result_json is empty")
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        from json_repair import repair_json

        fixed = repair_json(raw)
        return json.loads(fixed)


def tailoring_result_from_json(s: str) -> List[TailoredBlock]:
    data = _loads_tailoring_payload(s)
    if isinstance(data, dict) and "blocks" in data:
        data = data["blocks"]
    if not isinstance(data, list):
        raise ValueError("tailoring JSON must be a list or {\"blocks\": [...]}")
    return [TailoredBlock.model_validate(x) for x in data]


def read_docx_plain_text(docx_path: str) -> str:
    d = Document(docx_path)
    return "\n".join(p.text for p in _iter_docx_paragraphs(d))


def structured_resume_to_plain_text(content: StructuredResumeContent) -> str:
    """Flat text for evaluation / similarity (mirrors PDF sections)."""
    lines: List[str] = [
        content.full_name,
        content.headline,
        content.contact_line,
    ]
    if content.linkedin_url.strip():
        lines.append(content.linkedin_url)
    lines.append("")
    lines.append("PROFESSIONAL SUMMARY")
    lines.append(content.professional_summary)
    lines.append("")
    lines.append("PROFESSIONAL EXPERIENCE")
    for job in content.experience:
        lines.append(job.title)
        lines.append(f"{job.company}\t{job.date_range}")
        for b in job.bullets:
            lines.append(f"• {b}")
        if job.tech_used.strip():
            lines.append(f"Tech Used: {job.tech_used}")
        lines.append("")
    lines.append("TECHNICAL SKILLS")
    lines.append(f"Frontend | {content.technical_skills_frontend}")
    lines.append(f"Backend | {content.technical_skills_backend}")
    lines.append(f"DBs & DevOps | {content.technical_skills_devops}")
    lines.append("")
    lines.append("EDUCATION")
    for edu in content.education:
        lines.append(edu.institution)
        lines.append(f"{edu.degree_line}\t{edu.dates}")
    return "\n".join(lines).strip() + "\n"


def render_structured_resume_pdf(content: StructuredResumeContent, out_path: Path) -> None:
    """
    Build a PDF that follows the layout of input/resume_template.pdf:
    header, summary, experience with bullets and Tech Used, skill rows, education.
    """
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Paragraph as RLParagraph
    from reportlab.platypus import SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    name_style = ParagraphStyle(
        "TRName",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=17,
        spaceAfter=3,
    )
    headline_style = ParagraphStyle(
        "TRHeadline",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=12,
        spaceAfter=2,
    )
    contact_style = ParagraphStyle(
        "TRContact",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=11,
        spaceAfter=2,
    )
    section_style = ParagraphStyle(
        "TRSection",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=12,
        spaceBefore=10,
        spaceAfter=5,
    )
    body_style = ParagraphStyle(
        "TRBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=12,
        spaceAfter=6,
        alignment=TA_JUSTIFY,
    )
    role_title_style = ParagraphStyle(
        "TRRoleTitle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=12,
        spaceBefore=5,
        spaceAfter=2,
    )
    role_meta_style = ParagraphStyle(
        "TRRoleMeta",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=12,
        spaceAfter=4,
    )
    bullet_style = ParagraphStyle(
        "TRBullet",
        parent=body_style,
        leftIndent=12,
        firstLineIndent=-10,
        spaceAfter=3,
        bulletIndent=0,
    )
    tech_style = ParagraphStyle(
        "TRTech",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=9,
        leading=11,
        spaceAfter=6,
        alignment=TA_JUSTIFY,
    )

    story: List[Any] = []
    story.append(RLParagraph(_docx_text_to_reportlab_markup(content.full_name), name_style))
    if content.headline.strip():
        story.append(RLParagraph(_docx_text_to_reportlab_markup(content.headline), headline_style))
    if content.contact_line.strip():
        story.append(RLParagraph(_docx_text_to_reportlab_markup(content.contact_line), contact_style))
    if content.linkedin_url.strip():
        story.append(RLParagraph(_docx_text_to_reportlab_markup(content.linkedin_url), contact_style))
    story.append(Spacer(1, 4))

    story.append(RLParagraph("PROFESSIONAL SUMMARY", section_style))
    story.append(RLParagraph(_docx_text_to_reportlab_markup(content.professional_summary), body_style))

    story.append(RLParagraph("PROFESSIONAL EXPERIENCE", section_style))
    for job in content.experience:
        story.append(RLParagraph(_docx_text_to_reportlab_markup(job.title), role_title_style))
        meta = f"{job.company}\t{job.date_range}"
        story.append(RLParagraph(_docx_text_to_reportlab_markup(meta), role_meta_style))
        for bullet in job.bullets:
            story.append(
                RLParagraph("• " + _docx_text_to_reportlab_markup(bullet), bullet_style)
            )
        if job.tech_used.strip():
            tu = "Tech Used: " + job.tech_used
            story.append(RLParagraph(_docx_text_to_reportlab_markup(tu), tech_style))

    story.append(RLParagraph("TECHNICAL SKILLS", section_style))

    def _skill_line(label: str, body: str) -> str:
        b = (body or "").strip()
        lab = label.replace("&", "&amp;").replace("<", "&lt;")
        inner = _docx_text_to_reportlab_markup(b).replace("<br/>", " ")
        return f"<b>{lab}</b> | {inner}"

    story.append(
        RLParagraph(_skill_line("Frontend", content.technical_skills_frontend), body_style)
    )
    story.append(
        RLParagraph(_skill_line("Backend", content.technical_skills_backend), body_style)
    )
    story.append(
        RLParagraph(_skill_line("DBs & DevOps", content.technical_skills_devops), body_style)
    )

    story.append(RLParagraph("EDUCATION", section_style))
    for edu in content.education:
        story.append(RLParagraph(_docx_text_to_reportlab_markup(edu.institution), role_title_style))
        edu_line = f"{edu.degree_line}\t{edu.dates}"
        story.append(RLParagraph(_docx_text_to_reportlab_markup(edu_line), role_meta_style))

    pdf = SimpleDocTemplate(
        str(out_path),
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=48,
        bottomMargin=48,
        title="Resume",
    )
    pdf.build(story)


def loads_structured_resume_json(s: str) -> StructuredResumeContent:
    data = _loads_tailoring_payload(s)
    return StructuredResumeContent.model_validate(data)
